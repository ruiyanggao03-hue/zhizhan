import os
import json
import uuid
import tempfile
import re
from datetime import datetime
from typing import List, Dict, Optional
from pydantic import BaseModel
from fastapi import APIRouter, UploadFile, File, BackgroundTasks, Depends
from auth.auth_middleware import get_current_user
from auth.memory import get_memory_context
from fastapi.responses import StreamingResponse, FileResponse, JSONResponse
from dotenv import load_dotenv

from sentence_transformers import CrossEncoder
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.document_loaders import PyMuPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_openai import ChatOpenAI
from langchain_community.vectorstores import Chroma
from duckduckgo_search import DDGS

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv()

# ── 模块级"意图路由器" ──────────────────────────────────
# 只在文件被 Python 加载时创建一次，之后所有请求都复用这同一个实例，
# 避免每次意图分类都新建 ChatOpenAI（底层会重复创建 HTTP 连接池）。
_router_llm = ChatOpenAI(
    api_key=os.getenv("DEEPSEEK_API_KEY"),
    base_url="https://api.deepseek.com",
    model="deepseek-chat",
    temperature=0.0,          # 意图分类是"判断题"，温度最低保证输出稳定
    max_tokens=20,            # 只需要输出一个标签，20 个 token 足够
)

router = APIRouter(dependencies=[Depends(get_current_user)])

# =====================================================================
# 意图类型
# =====================================================================
INTENT_OFF_TOPIC = "off_topic"
INTENT_CLARIFY = "clarify"
INTENT_FOCUSED = "focused"
INTENT_REVISE = "revise"
INTENT_FULL_REPORT = "full_report"

VALID_INTENTS = {
    INTENT_OFF_TOPIC, INTENT_CLARIFY, INTENT_FOCUSED,
    INTENT_REVISE, INTENT_FULL_REPORT
}

class RagChatRequest(BaseModel):
    stock_code: str
    stock_name: str
    industry: str
    message: str
    selected_docs: List[str]
    history: List[Dict]
    conversation_id: str = ""  # 对话ID，用于加载记忆上下文

# =====================================================================
# 意图识别 & 研报校验
# =====================================================================
def _history_has_full_report(history: List[Dict]) -> bool:
    for msg in history:
        if msg.get("sender") == "ai" and "# 智瞻深度研报_" in (msg.get("text") or ""):
            return True
    return False

def classify_intent_by_rules(message: str, history: List[Dict]) -> Optional[str]:
    msg = message.strip()

    off_patterns = [
        "讲个笑话", "讲笑话", "写首诗", "写一首诗", "写代码",
        "天气怎么样", "今天天气", "你是谁", "你叫什么", "播放音乐", "唱首歌"
    ]
    if any(p in msg for p in off_patterns):
        return INTENT_OFF_TOPIC

    full_report_keywords = [
        "完整研报", "完整深度研报", "深度研报", "生成研报", "出具报告",
        "标准格式", "写一份报告", "出具深度", "开始撰写", "生成完整",
        "写一份完整", "按标准格式"
    ]
    if any(k in msg for k in full_report_keywords):
        return INTENT_FULL_REPORT

    if _history_has_full_report(history):
        revise_keywords = ["修改", "补充", "重写", "更新第", "摘要改", "改一下", "调整", "润色", "扩充"]
        if any(k in msg for k in revise_keywords):
            return INTENT_REVISE

    focused_keywords = [
        "只分析", "重点看", "重点分析", "帮我讲讲", "不要写全文",
        "单看", "专题", "单独分析", "不用完整", "简要分析"
    ]
    if any(k in msg for k in focused_keywords):
        return INTENT_FOCUSED

    clarify_patterns = ["分析一下", "写个报告", "写研报", "帮我分析", "研究一下", "看看怎么样"]
    if any(p in msg for p in clarify_patterns) and len(msg) <= 20:
        return INTENT_CLARIFY

    return None

async def classify_intent_by_llm(message: str, history: List[Dict], stock_name: str) -> str:
    try:
        history_text = ""
        for h in history[-4:]:
            role = "用户" if h.get("sender") == "user" else "助手"
            history_text += f"{role}: {(h.get('text') or '')[:200]}\n"

        prompt = f"""你是投研对话路由。根据用户最新消息和对话历史，只输出一个标签（不要任何解释）：
OFF_TOPIC | CLARIFY | FOCUSED | REVISE | FULL_REPORT

判断标准：
- OFF_TOPIC：与股票/宏观/行业/投资完全无关（闲聊、写诗、写代码等）
- CLARIFY：投研相关但范围不清，需先确认分析角度或框架
- FOCUSED：只要某一主题的深度分析，不要完整研报
- REVISE：要求修改、补充、调整已有研报（历史中已有完整研报时）
- FULL_REPORT：明确要求完整深度研报，或需求已充分明确且用户确认开始撰写

标的：{stock_name}
对话历史：
{history_text}
用户最新消息：{message}

只输出一个标签："""

        response = await _router_llm.ainvoke(prompt)
        label = response.content.strip().upper()
        mapping = {
            "OFF_TOPIC": INTENT_OFF_TOPIC,
            "CLARIFY": INTENT_CLARIFY,
            "FOCUSED": INTENT_FOCUSED,
            "REVISE": INTENT_REVISE,
            "FULL_REPORT": INTENT_FULL_REPORT,
        }
        for key, val in mapping.items():
            if key in label:
                return val
    except Exception:
        pass
    return INTENT_CLARIFY

async def classify_intent(message: str, history: List[Dict], stock_name: str) -> str:
    rule_result = classify_intent_by_rules(message, history)
    if rule_result:
        return rule_result
    return await classify_intent_by_llm(message, history, stock_name)

def validate_complete_report(text: str):
    """校验报告是否为可导出的完整深度研报（不绑定特定股票，只看格式）。
    返回 (is_valid: bool, reason: str) —— reason 用于诊断哪一项未通过。"""
    if not text or len(text) < 3000:
        return False, f"文本过短（{len(text) if text else 0} 字），需 ≥ 3000 字"

    # 只要以 "# 智瞻深度研报_" 开头，不管具体是哪个股票/行业都放行
    has_title = bool(re.search(r"# 智瞻深度研报_", text))
    has_toc = "## 目 录" in text or "## 目录" in text
    has_chapters = bool(re.search(r"## [一二三四五六七八九十]+、", text))

    checks = [
        ("标题", has_title, "报告中未找到 '# 智瞻深度研报_' 标题"),
        ("标的证券", "**标的证券**" in text, "报告中缺少 '**标的证券**' 字段"),
        ("摘要", "**摘要**" in text, "报告中缺少 '**摘要**' 字段"),
        ("中文章节", has_chapters, "报告中未找到 '## 一、' 格式的中文章节标题"),
    ]

    for name, passed, fail_msg in checks:
        if not passed:
            return False, fail_msg

    return True, "ok"


def _extract_stock_from_report(text: str) -> str:
    """从研报标题中提取股票名称，用于导出文件名。"""
    m = re.search(r"# 智瞻深度研报_(.+?)_", text)
    return m.group(1) if m else "标的"

# =====================================================================
# Prompt 构建器
# =====================================================================
def _base_context(req, current_date_day: str, retrieved_context: str) -> str:
    return f"""
当前物理时间：【{current_date_day}】。
本次研报核心标的：【{req.stock_name} ({req.stock_code})】。
所属宏观生态赛道：【{req.industry}】。

【用户本次核心诉求】（必须优先响应）
"{req.message}"

【参考数据】
{retrieved_context}
"""

def build_off_topic_prompt(req, current_date_day: str) -> str:
    return f"""
你是智瞻首席行业研究员 Ruiyang，专业、简洁、有温度但不刻板。
当前标的：【{req.stock_name} ({req.stock_code})】，行业：【{req.industry}】。

用户提问与投研无关。请用 50-150 字：
1. 礼貌简短说明你的职责边界
2. 引导用户回到对该标的的行业/投资分析
严禁输出完整研报结构，严禁输出 # 智瞻深度研报_ 标题。
"""

def build_clarify_prompt(req, current_date_day: str, retrieved_context: str) -> str:
    return f"""
你是智瞻首席行业研究员 Ruiyang。专业但不刻板，像资深分析师与客户沟通。
{_base_context(req, current_date_day, retrieved_context)}

【当前任务：需求对齐，不写完整研报】
1. 用 1-2 句话复述你理解的用户关注点
2. 若需求仍模糊，提出 1-2 个精准追问，或给出 2-3 个可选分析框架供选择
3. 可基于参考数据做简要专业回应（200-600 字），但禁止输出完整研报
4. 结尾自然引导：如需完整深度研报，请明确说「生成完整深度研报」

严禁：
- 输出 # 智瞻深度研报_ 标题
- 输出摘要/目录/结语等完整研报结构
- 使用无序列表或有序列表
"""

def build_focused_prompt(req, current_date_day: str, retrieved_context: str) -> str:
    return f"""
你是智瞻首席行业研究员 Ruiyang。
{_base_context(req, current_date_day, retrieved_context)}

【当前任务：单主题深度分析】
围绕用户诉求「{req.message}」展开 800-2000 字专业分析。

规则：
1. 可用 ## / ### 小标题组织，但不要目录、不要摘要块、不要结语与展望
2. 禁止 # 智瞻深度研报_ 开头
3. 论述写成完整学术段落，禁止列表
4. 可插入 1-2 个 Markdown 表格辅助说明
5. 严禁买入/卖出建议；保持客观中立
"""

def build_revise_prompt(req, current_date_day: str, retrieved_context: str, last_report: str) -> str:
    return f"""
你是智瞻首席行业研究员 Ruiyang。
{_base_context(req, current_date_day, retrieved_context)}

【当前任务：修改已有研报】
用户要求：{req.message}

【已有研报内容（供修改参考）】
{last_report[:8000]}

规则：
1. 只修改用户要求的部分，不要无关重写
2. 若用户要求的是局部调整，输出修改后的相关章节即可，不要强制全文
3. 若用户明确要求更新全文且输出完整结构，才使用完整研报模板
4. 局部修改时禁止 # 智瞻深度研报_ 开头
"""
def build_full_report_prompt(req, current_date_day: str, industry_main: str, retrieved_context: str) -> str:
    # ===== 完整研报 Prompt：保持原有框架不变 =====
    return f"""
你现在是全球顶尖券商的"首席行业研究员"（Head of Industry Research）。
当前物理时间：【{current_date_day}】。
本次研报核心标的：【{req.stock_name} ({req.stock_code})】。
所属宏观生态赛道：【{req.industry}】。

【用户本次核心诉求】（撰写时必须融入章节选择与论述重心）
"{req.message}"

【交互与响应绝对原则】（系统最高指令，优先于其他所有规则）：
1. ⛔ 专注垂类，拒绝闲聊：你是一个高度专业的金融智库系统。如果用户的提问与股票、宏观经济、行业生态或投资完全无关（例如日常闲聊、写诗、写代码、讲笑话等），请以首席分析师的口吻礼貌且简短地拒绝，并迅速引导用户将话题拉回对【{req.stock_name}】的投研分析上，严禁生成长篇大论的非金融内容。
2. 🎯 高度响应用户需求：用户的具体提问是你本次撰写研报的"核心导向"。你必须深刻理解用户输入的特殊关注点，并将用户关心的重点问题【直接转化为研报的核心分析章节】。绝对不可自说自话、忽略用户的显性指令。

【研报核心定调与研究框架】：
1. 拒绝千篇一律的AI生成感！这是一份针对【{req.industry}】量身定制的【行业深度与竞争格局】研报。
2. 必须采用"自上而下（Top-Down）"的华尔街研究框架：但【不必拘泥于固定的章节名称】。请根据该行业的特殊性自由且动态地展开论述。
3. 严禁通篇只围绕个股进行预测，严禁给出"买入/卖出/建仓"等低级直接投资建议！分析的终点应落脚于"商业模式演进与尾部风险定价"。

【研报结构与格式绝对纪律】：
1. 🌟 强制开头格式（原封不动输出）：
# 智瞻深度研报_{req.stock_name}_{industry_main}

**标的证券**：{req.stock_name}
**所属行业**：{req.industry}
**当前时间**：{current_date_day}
**分析师**：研究员Ruiyang。

**摘要**：
（高度凝练300字左右：一句话定调行业核心矛盾，一句话总结竞争格局，一句话透视标的核心壁垒，一句话指出最大潜在风险）

---

2. 🌟 灵活且极具深度的动态大纲（紧扣用户需求）：
   在上述内容之后，必须依次包含：
   - **## 目 录**（自动生成包含所有级别的纯文本大纲树，严禁加虚线或伪页码）
   - **正文章节**：根据【{req.industry}】的行业特征，【结合用户刚才提出的具体提问】，动态拟定 5 到 8 个极具专业度的一级章节标题（例如不要用干瘪的"宏观环境"，而是用"一、AI算力爆发下的产业周期重构"）。
     * 💡 如果用户有特定的关切（如"帮我重点分析它的海外业务"或"分析它的毛利率下滑"），必须将其直接作为独立的核心章节进行超深度拆解！
   - **## 结语与展望**（作为最后一个一级标题总结全篇）

3. 🌟 标题编号语法（严格遵守，用于系统排版引擎精准分页）：
   - 【一级标题】：Markdown 二级，必须以中文大写数字开头（如：## 一、 算力基建产业周期演进）。
   - 【二级标题】：Markdown 三级（如：### 1.1 核心组件市场规模）。
   - 【三级标题】：Markdown 四级（如：#### 1.1.1 需求端核心驱动因素）。（严禁在正文中使用单 # 号）

4. ⛔ 【排版格式高压红线】（绝对禁止零碎列表，确保公文质感）：
   - 绝对禁止使用任何无序列表（- 或 *）或有序列表（1. 2. 3.）！
   - 绝对禁止使用"短句：解释"这种聊天体结构！
   - 所有的论述、观点、数据，必须写成【完整的、逻辑严密的学术大段落】！段落间通过过渡词自然衔接。

5. 🌟 图表插入哲学（自然、合适、不生搬硬套）：
   - 寻找【最适合进行数据对比、财务指标拆解、竞品特征罗列】的地方，自然地插入 3~4 个 Markdown 表格。不要生搬硬套，只要做到整体图文并茂、数据清晰即可。
   - 表格上方必须紧跟居中标题，格式为：**表1 XXXXX对比表**。

【内容撰写深度与数据高压纪律】：
1. 坚守真实底线：数据和内容不能凭空捏造，必须严格依托底部的客观【参考数据】。严禁出现 AI 幻觉！若缺乏精确数据，请基于定性逻辑进行推演，切勿编造虚假财报。
2. 数据融合权重：优先挖掘【知识库资料】构建行业底层框架与护城河逻辑；同时高度重视【全网实时动态补充】，将其作为近1-3年最新政策、突发事件的时效性支撑，不可忽略。
3. 极度客观中立：严禁主观吹捧。深刻剖析该行业及标的面临的【政策风险、技术挑战和竞争劣势】等残酷风险。
4. 首席分析师语境：绝对自信且专业，不要输出类似免责声明的话（如"股市有风险"、"作为AI我无法预测..."）。
5. 总字数必须在 6000 字以上（可达12000字），做到事无巨细、穿透底层逻辑，篇幅要极长且详实！

【参考数据】
{retrieved_context}
"""

def build_system_prompt(intent: str, req, current_date_day: str, industry_main: str, retrieved_context: str, history: List[Dict], memory_ctx: dict = None) -> str:
    # Build the base prompt for this intent
    if intent == INTENT_OFF_TOPIC:
        base = build_off_topic_prompt(req, current_date_day)
    elif intent == INTENT_CLARIFY:
        base = build_clarify_prompt(req, current_date_day, retrieved_context)
    elif intent == INTENT_FOCUSED:
        base = build_focused_prompt(req, current_date_day, retrieved_context)
    elif intent == INTENT_REVISE:
        last_report = ""
        for msg in reversed(history):
            if msg.get("sender") == "ai" and "# 智瞻深度研报_" in (msg.get("text") or ""):
                last_report = msg["text"]
                break
        base = build_revise_prompt(req, current_date_day, retrieved_context, last_report)
    else:
        base = build_full_report_prompt(req, current_date_day, industry_main, retrieved_context)

    # Inject conversation memory if available
    if memory_ctx:
        memory_block = ""
        if memory_ctx.get("key_facts_text"):
            memory_block += f"【本对话已提炼的关键事实】\n{memory_ctx['key_facts_text']}\n\n"
        if memory_ctx.get("summary_text"):
            memory_block += f"【早期对话摘要】\n{memory_ctx['summary_text']}\n\n"
        if memory_block:
            base = f"【对话记忆】（这是你与用户之前的对话要点，请记住并主动引用这些信息）\n{memory_block}\n{base}"

    return base

# =====================================================================
# RAG 引擎
# =====================================================================
class ZhiZhanRAGEngine:
    def __init__(self):
        deepseek_key = os.getenv("DEEPSEEK_API_KEY")
        if not deepseek_key:
            raise ValueError("❌ 未在 .env 中找到 DEEPSEEK_API_KEY")

        self.llm = ChatOpenAI(
            api_key=deepseek_key,
            base_url="https://api.deepseek.com",
            model="deepseek-chat",
            temperature=0.4,
            max_tokens=8192,
            streaming=True
        )
        self.embeddings = HuggingFaceEmbeddings(
            model_name="BAAI/bge-small-zh-v1.5",
            model_kwargs={'device': 'cpu'},
            encode_kwargs={'normalize_embeddings': True}
        )
        self.persist_directory = "./chroma_db"

        self.system_db = Chroma(
            collection_name="zhizhan_system_reports",
            embedding_function=self.embeddings,
            persist_directory=self.persist_directory
        )
        self.private_db = Chroma(
            collection_name="zhizhan_ephemeral_reports",
            embedding_function=self.embeddings
        )
        # CPU 友好：max_length=64，判断文档相关性无需长文本
        self.reranker = CrossEncoder('BAAI/bge-reranker-v2-m3', max_length=64)
        self._industry_cache = {}

    async def get_industry_by_llm(self, stock_code: str, stock_name: str) -> str:
        # 行业分类不会变，用内存缓存避免重复调用 AI
        cache_key = str(stock_code).strip()
        if cache_key in self._industry_cache:
            return self._industry_cache[cache_key]

        try:
            prompt = f"请告诉我上市公司【{stock_name} (代码:{stock_code})】所属的【宏观大行业-细分核心赛道】。只输出这几个字，严禁输出任何其他解释或标点符号。"
            response = await self.llm.ainvoke(prompt)
            result = response.content.strip()
            self._industry_cache[cache_key] = result
            return result
        except Exception:
            return "全市场综合概况"

    def fetch_web_news(self, query: str) -> str:
        try:
            from concurrent.futures import ThreadPoolExecutor, TimeoutError as FutureTimeout
            def _search():
                with DDGS() as ddgs:
                    current_year = datetime.now().year
                    results = list(ddgs.text(
                        f"{query} {current_year}最新行业概况 深度研报 数据 进展 风险挑战",
                        max_results=3
                    ))
                    return "\n".join([
                        f"[全网动态补充 {i+1}]: {r['title']} - {r['body']}"
                        for i, r in enumerate(results)
                    ])
            with ThreadPoolExecutor(max_workers=1) as executor:
                future = executor.submit(_search)
                return future.result(timeout=8)
        except (Exception, FutureTimeout):
            return "全网最新资讯加载超时。"

    def delete_private_doc(self, doc_id: str):
        try:
            self.private_db._collection.delete(where={"doc_id": doc_id})
        except Exception:
            pass

    async def generate_rag_stream(self, req: RagChatRequest):
        intent = await classify_intent(req.message, req.history, req.stock_name)

        retrieved_context = ""
        raw_docs = []

        try:
            raw_docs.extend(self.system_db.similarity_search(req.message, k=6))
        except Exception:
            pass

        if req.selected_docs:
            try:
                raw_docs.extend(self.private_db.similarity_search(
                    req.message, k=4, filter={"doc_id": {"$in": req.selected_docs}}
                ))
            except Exception:
                pass

        if raw_docs:
            # 去重
            seen = set()
            unique_docs = []
            for doc in raw_docs:
                if doc.page_content not in seen:
                    seen.add(doc.page_content)
                    unique_docs.append(doc)

            # 轻量 rerank：max_length=128，文档少速度快
            sentence_pairs = [[req.message, doc.page_content] for doc in unique_docs]
            rerank_scores = self.reranker.predict(sentence_pairs)
            scored_docs = list(zip(unique_docs, rerank_scores))
            scored_docs.sort(key=lambda x: x[1], reverse=True)
            top_docs = [doc for doc, score in scored_docs[:5] if score > 0]

            for idx, doc in enumerate(top_docs):
                source_name = doc.metadata.get('source', '内部研报')
                retrieved_context += f"[知识库资料 {idx+1} ({source_name})]: {doc.page_content}\n\n"

        retrieved_context += f"\n【全网实时动态补充】:\n{self.fetch_web_news(req.message + ' ' + req.industry)}\n"

        if not retrieved_context.strip():
            retrieved_context = "请基于您的专业知识储备进行深度推演扩写。"

        current_date_day = datetime.now().strftime("%Y年%m月%d日")
        industry_main = req.industry.split('-')[0] if '-' in req.industry else req.industry

        # Load conversation memory
        memory_ctx = None
        if req.conversation_id:
            try:
                memory_ctx = get_memory_context(req.conversation_id)
            except Exception:
                pass

        system_prompt = build_system_prompt(
            intent, req, current_date_day, industry_main, retrieved_context, req.history, memory_ctx
        )

        messages = [{"role": "system", "content": system_prompt}]
        for msg in req.history:
            if msg.get("text"):
                role = "user" if msg["sender"] == "user" else "assistant"
                messages.append({"role": role, "content": msg["text"]})
        messages.append({"role": "user", "content": req.message})

        yield f"data: {json.dumps({'type': 'meta', 'intent': intent, 'exportable': False}, ensure_ascii=False)}\n\n"

        full_content = ""
        async for chunk in self.llm.astream(messages):
            if chunk.content:
                full_content += chunk.content
                yield f"data: {json.dumps({'type': 'answer', 'content': chunk.content}, ensure_ascii=False)}\n\n"

        is_valid, _ = validate_complete_report(full_content)
        exportable = (
            intent == INTENT_FULL_REPORT and is_valid
        )
        yield f"data: {json.dumps({'type': 'done', 'intent': intent, 'exportable': exportable}, ensure_ascii=False)}\n\n"
        yield "data: [DONE]\n\n"

# =====================================================================
# 工具函数区（Word 导出，保持原样）
# =====================================================================
def set_run_font(run, ascii_font='Times New Roman', east_asia_font='宋体'):
    run.font.name = ascii_font
    run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)

def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def generate_native_docx(md_text: str, file_path: str):
    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21.0), Cm(29.7)
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Cm(2.5)

    footer = section.footer
    p_footer = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p_footer.add_run("- "), 'Times New Roman', '宋体')
    add_page_number(p_footer.add_run())
    set_run_font(p_footer.add_run(" -"), 'Times New Roman', '宋体')

    lines = md_text.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line:
            i += 1
            continue

        if line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i].strip())
                i += 1

            raw_rows = []
            for tl in table_lines:
                if '---' in tl:
                    continue
                parts = [p.strip() for p in tl.split('|')]
                if len(parts) >= 2:
                    raw_rows.append(parts[1:-1])

            if raw_rows:
                num_rows = len(raw_rows)
                num_cols = max(len(r) for r in raw_rows)
                table = doc.add_table(rows=num_rows, cols=num_cols)
                table.style = 'Table Grid'
                for r_idx, row_data in enumerate(raw_rows):
                    for c_idx, cell_value in enumerate(row_data):
                        if c_idx < num_cols:
                            cell = table.cell(r_idx, c_idx)
                            cell.text = cell_value.replace('**', '')
                            p_cell = cell.paragraphs[0]
                            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p_cell.paragraph_format.space_before = Pt(4)
                            p_cell.paragraph_format.space_after = Pt(4)
                            for run in p_cell.runs:
                                if r_idx == 0:
                                    run.bold = True
                                    set_run_font(run, 'SimHei', '黑体')
                                else:
                                    set_run_font(run, 'Times New Roman', '宋体')
                                run.font.size = Pt(10.5)
            continue

        p = doc.add_paragraph()

        if line.startswith('# '):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.space_after = Pt(24)
            run = p.add_run(line[2:].strip())
            set_run_font(run, 'SimHei', '黑体')
            run.font.size = Pt(26)
            run.bold = True

        elif line.startswith('## '):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.space_before = Pt(22)
            p.paragraph_format.space_after = Pt(22)
            p.paragraph_format.page_break_before = True
            run = p.add_run(line[3:].strip())
            set_run_font(run, 'SimHei', '黑体')
            run.font.size = Pt(22)
            run.bold = True

        elif line.startswith('### '):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.space_before = Pt(10)
            p.paragraph_format.space_after = Pt(10)
            run = p.add_run(line[4:].strip())
            set_run_font(run, 'SimHei', '黑体')
            run.font.size = Pt(16)
            run.bold = True

        elif line.startswith('#### '):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.space_before = Pt(10)
            run = p.add_run(line[5:].strip())
            set_run_font(run, 'SimHei', '黑体')
            run.font.size = Pt(14)
            run.bold = True

        elif line.startswith('**表') or line.startswith('**图') or line.startswith('**目录') or line.startswith('**摘要'):
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.5
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = 0
            run = p.add_run(line.replace('**', ''))
            set_run_font(run, 'SimHei', '黑体')
            run.font.size = Pt(10.5)
            run.bold = True

        else:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            p.paragraph_format.line_spacing = Pt(20)
            p.paragraph_format.first_line_indent = Cm(0.85)

            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    set_run_font(run, 'Times New Roman', '宋体')
                    run.font.size = Pt(12)
                else:
                    run = p.add_run(part)
                    set_run_font(run, 'Times New Roman', '宋体')
                    run.font.size = Pt(12)
        i += 1
    doc.save(file_path)

# =====================================================================
# 接口路由
# =====================================================================
rag_engine = ZhiZhanRAGEngine()

def cleanup_temp_files(*file_paths):
    for path in file_paths:
        try:
            if os.path.exists(path):
                os.remove(path)
                print(f"🧹 内存释放：已成功清理临时文件 {path}")
        except Exception as e:
            print(f"⚠️ 清理临时文件失败 {path}: {e}")

@router.get("/api/industry")
async def get_industry_endpoint(code: str, name: str = ""):
    clean_code = str(code).strip()
    result = await rag_engine.get_industry_by_llm(clean_code, name)
    return {"status": "success", "industry": result}

@router.post("/api/rag/chat")
async def rag_chat_endpoint(request: RagChatRequest):
    generator = rag_engine.generate_rag_stream(request)
    return StreamingResponse(generator, media_type="text/event-stream")

@router.post("/api/rag/upload")
async def upload_private_document(file: UploadFile = File(...)):
    doc_id = str(uuid.uuid4())
    with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        content = await file.read()
        tmp.write(content)
        tmp_path = tmp.name

    try:
        loader = PyMuPDFLoader(tmp_path)
        docs = loader.load()
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=800, chunk_overlap=150)
        chunks = text_splitter.split_documents(docs)
        for chunk in chunks:
            chunk.metadata["doc_id"] = doc_id
            chunk.metadata["source"] = file.filename
        rag_engine.private_db.add_documents(chunks)
        return {"status": "success", "doc_id": doc_id, "title": file.filename}
    except Exception as e:
        return {"status": "error", "message": str(e)}
    finally:
        if os.path.exists(tmp_path):
            os.remove(tmp_path)

class DeleteDocRequest(BaseModel):
    doc_id: str

@router.post("/api/rag/delete_doc")
async def delete_private_document(req: DeleteDocRequest):
    rag_engine.delete_private_doc(req.doc_id)
    return {"status": "success", "message": "内存已释放"}

class ExportRequest(BaseModel):
    stock_name: str
    industry: str = ""
    markdown_content: str

@router.post("/api/rag/export/{format}")
async def export_report(format: str, req: ExportRequest, background_tasks: BackgroundTasks):
    is_valid, fail_reason = validate_complete_report(req.markdown_content)
    if not is_valid:
        print(f"[导出失败] reason={fail_reason}")
        print(f"[导出失败] 报告前200字: {(req.markdown_content or '')[:200]}")
        return JSONResponse(
            status_code=400,
            content={"status": "error", "message": f"报告格式校验未通过：{fail_reason}。请明确要求「生成完整深度研报」后重试。"}
        )

    # 从报告内容中提取股票名（而非依赖前端当前页面），历史报告也能正确命名
    report_stock_name = _extract_stock_from_report(req.markdown_content)

    tmp_docx = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    tmp_docx_path = tmp_docx.name
    tmp_docx.close()

    generate_native_docx(req.markdown_content, tmp_docx_path)

    if format == 'word':
        background_tasks.add_task(cleanup_temp_files, tmp_docx_path)
        return FileResponse(
            tmp_docx_path,
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            filename=f"智瞻深度研报_{report_stock_name}.docx"
        )

    elif format == 'pdf':
        import subprocess
        tmp_pdf_path = None
        try:
            tmp_pdf_expected = tmp_docx_path.replace('.docx', '.pdf')

            # 使用 LibreOffice 无头模式转换（Ubuntu 服务器兼容）
            result = subprocess.run(
                ['libreoffice', '--headless', '--convert-to', 'pdf', '--outdir',
                 os.path.dirname(tmp_docx_path), tmp_docx_path],
                capture_output=True, text=True, timeout=60
            )

            if result.returncode != 0 or not os.path.exists(tmp_pdf_expected):
                raise RuntimeError(result.stderr or 'LibreOffice conversion failed')

            tmp_pdf_path = tmp_pdf_expected

            background_tasks.add_task(cleanup_temp_files, tmp_docx_path, tmp_pdf_path)
            return FileResponse(
                tmp_pdf_path,
                media_type='application/pdf',
                filename=f"智瞻深度研报_{report_stock_name}.pdf"
            )

        except Exception as e:
            cleanup_temp_files(tmp_docx_path)
            if tmp_pdf_path:
                cleanup_temp_files(tmp_pdf_path)
            return JSONResponse(
                status_code=500,
                content={"status": "error", "message": f"PDF导出失败：{str(e)}"}
            )

    return JSONResponse(status_code=400, content={"status": "error", "message": "不支持的导出格式"})